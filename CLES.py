import sys
import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from docx import Document
from docx.shared import Inches
from PyQt5.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,QHBoxLayout, QTabWidget, QLabel, QPushButton, QLineEdit, QTextEdit, QComboBox, QFrame, QSplitter,QScrollArea, QMessageBox, QFileDialog, QProgressBar,QStatusBar, QGroupBox, QGridLayout, QTableWidget,QTableWidgetItem, QHeaderView, QAbstractItemView)
from PyQt5.QtGui import QFont, QPalette, QColor, QPixmap, QIcon, QPainter, QLinearGradient
from PyQt5.QtCore import Qt, QTimer, QPropertyAnimation, QEasingCurve, QRect
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
plt.rcParams['font.sans-serif'] = ['SimHei'] 
plt.rcParams['axes.unicode_minus'] = False
# ====================== 模块一：数据管理逻辑 ======================
class CLESDataManager:
    def __init__(self, file_path):
        self.file_path = file_path
        self.df = pd.read_excel(file_path, sheet_name='Sheet1')
        
    def save(self):
        try:
            self.df.to_excel(self.file_path, index=False, sheet_name='Sheet1')
            return True
        except Exception as e:
            QMessageBox.warning(None, "保存失败", f"文件被占用，请关闭Excel后重试：\n{str(e)}")
            return False

# ====================== 模块二：增强型统计分析引擎 ======================
class CLESStatsEngine:
    def __init__(self, df):
        self.df = df
        self.cols = df.columns.tolist()

    def _get_val(self, col, func='sum'):
        if col in self.cols:
            if func == 'sum': return self.df[col].sum()
            if func == 'mean': return self.df[col].mean()
            if func == 'max': return self.df[col].max()
        return 0

    def run_full_analysis(self):
        df = self.df
        cols = self.cols
        res = []

        # 头部装饰
        res.append("╔" + "═"*60 + "╗")
        res.append(f"║{'农村土地经济调研数据分析系统V1.0 深度描述性分析报告':^52}║")
        res.append(f"║{'生成时间: ' + pd.Timestamp.now().strftime('%Y-%m-%d %H:%M'):^56}║")
        res.append("╚" + "═"*60 + "╝\n")

        # --- 基础元数据统计 ---
        valid_n = len(df)
        res.append(f"【 📊 样本概况 】")
        res.append(f" • 有效观测值: {valid_n} 户")
        if 'a104' in cols:
            res.append(f" • 调查区域分布: 共涉及 {df['a104'].nunique()} 个行政村")
        res.append("="*62)

        # --- 一、家庭禀赋与人口结构 (A模块) ---
        res.append("\n【 👤 一、家庭禀赋与劳动力结构 】")
        if 'a101' in cols:
            stats = df['a101'].describe()
            res.append(f" • 家庭规模: 均值 {stats['mean']:.2f}人 | 中位数 {stats['50%']:.0f}人 | 标准差 {stats['std']:.2f}")
        
        # 劳动力抚养比 (假设 a101是总人口, a108/a109是老人小孩)
        res.append(f" • 政治资本: ")
        if 'a102' in cols:
            res.append(f"   - 干部背景率: {df['a102'].mean():.1%} (反映社会资本能力)")
        if 'a103' in cols:
            res.append(f"   - 党员覆盖率: {df['a103'].mean():.1%} (反映政策响应潜力)")
        
        if 'a201a1' in cols: # 户主受教育年限
            edu_mean = df['a201a1'].mean()
            res.append(f" • 户主人力资本: 平均受教育年限 {edu_mean:.1f} 年")

        # --- 二、收入质量与多元化深度分析 (B模块) ---
        res.append("\n【 💰 二、家庭收入结构与稳定性分析 】")
        inc_cols = {
            '工资性收入(外出务工)': 'b101', 
            '经营性收入(农林牧渔)': 'b102', 
            '财产性收入(租金股息)': 'b103', 
            '转移性收入(政府补贴)': 'b104'
        }
        
        active_inc = {}
        for label, c in inc_cols.items():
            if c in cols: active_inc[label] = df[c].replace(np.nan, 0)

        if active_inc:
            total_inc = sum(active_inc.values())
            total_sum = total_inc.sum()
            avg_inc = total_inc.mean()
            
            res.append(f" • 收入总量: 样本户均总收入 ¥ {avg_inc:,.2f}")
            res.append(f" • 收入极差: 最大值 ¥ {total_inc.max():,.2f} | 最小值 ¥ {total_inc.min():,.2f}")
            
            # 集中度分析 (Herfindahl-Hirschman Index 变体)
            for label, series in active_inc.items():
                share = (series.sum() / total_sum * 100)
                # 计算该项收入的覆盖率
                participation = (series > 0).mean()
                res.append(f"   └─ {label}: 占比 {share:.1f}% | 农户参与率 {participation:.1%}")
            
            # 基尼系数估算 (简化版)
            sorted_inc = np.sort(total_inc)
            index = np.arange(1, len(sorted_inc)+1)
            gini = (np.sum((2 * index - len(sorted_inc) - 1) * sorted_inc)) / (len(sorted_inc) * np.sum(sorted_inc))
            res.append(f" • 收入分配公平度: 样本内基尼系数预估为 {gini:.3f}")

        # --- 三、土地资源配置与细碎化 (C模块) ---
        res.append("\n【 🌾 三、土地利用效率与流转倾向 】")
        land_cols = {'承包地': 'c101', '经营地': 'c102', '转入': 'c103', '转出': 'c104'}
        
        for label, c in land_cols.items():
            if c in cols:
                res.append(f" • {label}面积: 户均 {df[c].mean():.2f} 亩 | 样本总计 {df[c].sum():,.1f} 亩")
        
        if 'c103' in cols and 'c104' in cols:
            market_active = ((df['c103'] > 0) | (df['c104'] > 0)).mean()
            res.append(f" • 土地市场活跃度: {market_active:.1%} 的农户参与了土地流转")

        if 'c201' in cols and 'c202' in cols:
            total_land = df['c201'] + df['c202']
            if total_land.sum() > 0:
                paddy_ratio = df['c201'].sum() / total_land.sum()
                res.append(f" • 土地质量结构: 水田占比 {paddy_ratio:.1%} | 旱地占比 {1-paddy_ratio:.1%}")

        # --- 四、住房、生产与社会保障 (C3, E, G模块) ---
        res.append("\n【 🏡 四、居住资产、农业生产与保障安全网 】")
        if 'c302' in cols:
            res.append(f" • 住房财富: 平均自估值 ¥ {df['c302'].mean():,.0f} (财富效应参考指标)")
        
        if 'e101' in cols and 'e102' in cols:
            # 土地产出率
            yield_rate = df['e102'].sum() / df['e101'].replace(0, np.nan).sum()
            res.append(f" • 农业生产率: 平均亩产规模 {yield_rate:.2f} (单位/亩)")

        if 'f101' in cols:
            credit_access = (df['f101'] == 1).mean()
            res.append(f" • 金融获得性: 农户正规信贷获得率为 {credit_access:.1%}")

        if 'g101' in cols:
            pension_rate = (df['g101'] == 1).mean()
            res.append(f" • 社保保障水平: 城乡居民养老保险参保率 {pension_rate:.1%}")

        # 结论性尾部
        res.append("\n" + "═"*62)
        res.append("【 💡 分析提示 】")
        res.append(" 1. 均值受极端值影响较大，建议对比中位数。")
        res.append(" 2. 若基尼系数 > 0.4，需关注该样本区域内的贫富差距风险。")
        res.append(" 3. 参与率低而占比高的收入项，往往是决定农户贫富差距的关键。")
        res.append("═"*62)
        res.append(f"{'报告分析结束':^60}")
        
        return "\n".join(res)

# ====================== 模块三：可视化与报表生成模块 ======================
class CLESVisualizer:
    def __init__(self, manager):
        self.manager = manager
        self.df = manager.df

    def plot_land_distribution(self, canvas_frame):
        """地块面积分布柱状图 (基于C1-07经营总面积)"""
        if 'c107' not in self.df.columns: return
        
        plt.figure(figsize=(5, 4))
        self.df['c107'].hist(bins=20, color='skyblue', edgecolor='black')
        plt.title("农户经营土地面积分布 (亩)")
        plt.xlabel("面积")
        plt.ylabel("户数")
        
        self._render_to_tkinter(canvas_frame)

    def plot_income_structure(self, canvas_frame):
        """收入结构饼图 (基于B1/B2模块)"""
        inc_map = {'工资性': 'b101', '经营性': 'b102', '财产性': 'b103', '转移性': 'b104'}
        data = {k: self.df[v].sum() for k, v in inc_map.items() if v in self.df.columns}
        
        if not data: return
        plt.figure(figsize=(5, 4))
        plt.pie(data.values(), labels=data.keys(), autopct='%1.1f%%', startangle=140)
        plt.title("样本户均收入结构占比")
        
        self._render_to_tkinter(canvas_frame)

    def _render_to_tkinter(self, frame):
        for widget in frame.winfo_children():
            widget.destroy()
        canvas = FigureCanvasTkAgg(plt.gcf(), master=frame)
        canvas.draw()
        canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)

class CLESReportGenerator:
    def __init__(self, manager):
        self.manager = manager

    def export_word_report(self):
        """导出农村土地经济调研数据分析系统V1.0官方格式汇总报表"""
        doc = Document()
        doc.add_heading('农村土地经济调研数据分析系统V1.0 农户调查分析报告', 0)

        # 1. 基础汇总表
        doc.add_heading('一、样本基本信息汇总', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '指标'
        hdr_cells[1].text = '样本量'
        hdr_cells[2].text = '均值'
        hdr_cells[3].text = '最大值'

        metrics = [('家庭人口(a101)', 'a101'), ('经营面积(c107)', 'c107'), ('总支出(h201)', 'h201')]
        for label, col in metrics:
            if col in self.manager.df.columns:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = str(len(self.manager.df))
                row_cells[2].text = f"{self.manager.df[col].mean():.2f}"
                row_cells[3].text = f"{self.manager.df[col].max():.2f}"

        # 2. 地块明细表 (D模块)
        doc.add_heading('二、最大地块明细 (D1模块)', level=1)
        if 'd102' in self.manager.df.columns:
            detail_df = self.manager.df[['hid', 'd102', 'd108', 'd109']].head(10) # 仅展示前10条
            t2 = doc.add_table(rows=1, cols=4)
            t2.style = 'Table Grid'
            for i, col_name in enumerate(['农户ID', '面积(亩)', '灌溉能力', '肥力']):
                t2.rows[0].cells[i].text = col_name
            for _, row in detail_df.iterrows():
                row_cells = t2.add_row().cells
                for i, val in enumerate(row): row_cells[i].text = str(val)

        file_path = filedialog.asksaveasfilename(defaultextension=".docx")
        if file_path:
            doc.save(file_path)
            messagebox.showinfo("成功", f"报表已保存至: {file_path}")

# ====================== 模块四：自定义查询界面 ======================
class CLESAnyalyzerGUI:
    def __init__(self, root, manager):
        self.manager = manager
        self.viz = CLESVisualizer(manager)
        self.rep = CLESReportGenerator(manager)
        
        # 界面布局
        filter_frame = ttk.LabelFrame(root, text="自定义筛选查询")
        filter_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(filter_frame, text="村编码(a104):").grid(row=0, column=0)
        self.ent_village = ttk.Entry(filter_frame)
        self.ent_village.grid(row=0, column=1)
        
        ttk.Button(filter_frame, text="执行筛选", command=self.apply_filter).grid(row=0, column=2, padx=5)
        ttk.Button(filter_frame, text="导出Word报表", command=self.rep.export_word_report).grid(row=0, column=3)

        self.chart_frame = ttk.Frame(root)
        self.chart_frame.pack(fill="both", expand=True)
        
        ttk.Button(root, text="查看收入结构图", command=lambda: self.viz.plot_income_structure(self.chart_frame)).pack(side="left", padx=20)
        ttk.Button(root, text="查看土地分布图", command=lambda: self.viz.plot_land_distribution(self.chart_frame)).pack(side="left")

    def apply_filter(self):
        v_code = self.ent_village.get()
        if v_code:
            try:
                # 转换类型匹配CSV中的数值
                self.manager.df = self.manager.df[self.manager.df['a104'] == int(v_code)]
                messagebox.showinfo("筛选成功", f"当前显示村代码 {v_code} 的数据，共 {len(self.manager.df)} 条")
            except:
                messagebox.showerror("错误", "请输入有效的数字村编码")

# ====================== 主界面 ======================
class CLESApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.manager = None
        self.viz = None
        self.rep = None
        
        self.initUI()
        self.applyModernStyle()

    def initUI(self):
        self.setWindowTitle("农村土地经济调研数据分析系统V1.0")
        self.setGeometry(100, 100, 1400, 900)
        self.setWindowIcon(QIcon())  # 可以设置图标
        
        # 创建中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 主布局
        main_layout = QVBoxLayout(central_widget)
        
        # 顶部标题栏
        self.createHeader(main_layout)
        
        # 主标签页
        self.tab_widget = QTabWidget()
        main_layout.addWidget(self.tab_widget)
        
        # 创建各个标签页
        self.createDataTab()
        self.createAnalysisTab()
        self.createVisualizationTab()
        self.createQueryTab()
        
        # 状态栏
        self.status_bar = self.statusBar()
        self.status_bar.showMessage("💡 就绪：请加载 Excel 文件")
        
        # 连接信号
        self.connectSignals()

    def createHeader(self, parent_layout):
        header_frame = QFrame()
        # 移除固定高度，让它自适应
        header_frame.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #667eea, stop:1 #764ba2);
                border-radius: 3px;
                margin: 2px;
            }
        """)
        parent_layout.addWidget(header_frame)
        
        header_layout = QHBoxLayout(header_frame)
        header_layout.setContentsMargins(5, 2, 5, 2)  # 进一步减少边距
        
        # Logo区域
        logo_label = QLabel("🌾")
        logo_label.setFont(QFont("Arial", 20))  # 减小字体
        logo_label.setStyleSheet("color: white;")
        header_layout.addWidget(logo_label)
        
        title_layout = QVBoxLayout()
        title_label = QLabel("农村土地经济调研数据分析系统V1.0")
        title_label.setFont(QFont("微软雅黑", 12, QFont.Bold))  # 减小字体
        title_label.setStyleSheet("color: white;")
        
        subtitle_label = QLabel("北京林业大学所有·专业数据分析工具")
        subtitle_label.setFont(QFont("微软雅黑", 7))  # 减小字体
        subtitle_label.setStyleSheet("color: #e8f4f8;")
        
        title_layout.addWidget(title_label)
        title_layout.addWidget(subtitle_label)
        header_layout.addLayout(title_layout)
        header_layout.addStretch()
        
        # 文件路径区域
        path_frame = QFrame()
        path_frame.setStyleSheet("""
            QFrame {
                background: rgba(255,255,255,0.1);
                border-radius: 2px;
                padding: 1px;
            }
        """)
        path_layout = QHBoxLayout(path_frame)
        path_layout.setContentsMargins(3, 1, 3, 1)
        
        path_icon = QLabel("📁")
        path_icon.setStyleSheet("color: white;")
        path_layout.addWidget(path_icon)
        
        path_label = QLabel("数据文件:")
        path_label.setStyleSheet("color: white; font-weight: bold; font-size: 9pt;")
        path_layout.addWidget(path_label)
        
        self.path_entry = QLineEdit()
        self.path_entry.setStyleSheet("""
            QLineEdit {
                background: rgba(255,255,255,0.2);
                color: white;
                border: none;
                padding: 2px;
                border-radius: 1px;
                font-size: 9pt;
            }
        """)
        path_layout.addWidget(self.path_entry)
        
        header_layout.addWidget(path_frame)
        
        # 按钮区域
        button_layout = QHBoxLayout()
        
        self.browse_btn = self.createStyledButton("📂 浏览文件", "#f39c12", "#e67e22")
        self.load_btn = self.createStyledButton("⚡ 加载数据", "#27ae60", "#2ecc71")
        
        button_layout.addWidget(self.browse_btn)
        button_layout.addWidget(self.load_btn)
        header_layout.addLayout(button_layout)

    def createStyledButton(self, text, bg_color, hover_color):
        btn = QPushButton(text)
        btn.setFont(QFont("微软雅黑", 8, QFont.Bold))  # 减小字体
        btn.setStyleSheet(f"""
            QPushButton {{
                background-color: {bg_color};
                color: white;
                border: none;
                padding: 6px 12px;
                border-radius: 3px;
                font-size: 8pt;
            }}
            QPushButton:hover {{
                background-color: {hover_color};
            }}
            QPushButton:pressed {{
                background-color: {bg_color};
                padding-top: 7px;
                padding-bottom: 5px;
            }}
        """)
        return btn

    def createDataTab(self):
        self.data_tab = QWidget()
        self.tab_widget.addTab(self.data_tab, "📝 数据录入与管理")
        
        layout = QVBoxLayout(self.data_tab)
        
        # 数据概览面板
        overview_group = QGroupBox("📈 数据概览")
        overview_group.setStyleSheet(self.getGroupBoxStyle())
        layout.addWidget(overview_group)
        
        overview_layout = QGridLayout(overview_group)
        
        # 统计卡片
        self.createStatCard(overview_layout, "👥 样本总量", "0", "#3498db", 0, 0)
        self.createStatCard(overview_layout, "🏘️ 村编码数", "0", "#2ecc71", 0, 1)
        self.createStatCard(overview_layout, "💰 平均收入", "¥0", "#9b59b6", 1, 0)
        self.createStatCard(overview_layout, "🌾 平均耕地", "0亩", "#e74c3c", 1, 1)
        
        # 数据操作面板 - 改为3列布局
        data_operations_layout = QHBoxLayout()
        layout.addLayout(data_operations_layout)
        
        # 列1：字段参照面板
        fields_group = QGroupBox("📋 字段对照表")
        fields_group.setStyleSheet(self.getGroupBoxStyle())
        fields_layout = QVBoxLayout(fields_group)
        self.fields_text = QTextEdit()
        self.fields_text.setFont(QFont("Consolas", 10))
        self.fields_text.setReadOnly(True)
        fields_layout.addWidget(self.fields_text)
        data_operations_layout.addWidget(fields_group)
        
        # 列2：手动录入新行
        input_group = QGroupBox("➕ 手动录入新行")
        input_group.setStyleSheet(self.getGroupBoxStyle())
        input_layout = QVBoxLayout(input_group)
        self.input_text = QTextEdit()
        input_layout.addWidget(self.input_text)
        
        self.add_btn = self.createStyledButton("📥 提交并同步文件", "#2ecc71", "#27ae60")
        input_layout.addWidget(self.add_btn)
        data_operations_layout.addWidget(input_group)
        
        # 列3：单元格修正
        edit_group = QGroupBox("🔧 单元格修正")
        edit_group.setStyleSheet(self.getGroupBoxStyle())
        edit_layout = QVBoxLayout(edit_group)
        
        # 农户ID
        hid_layout = QHBoxLayout()
        hid_layout.addWidget(QLabel("农户 HID:"))
        self.hid_entry = QLineEdit()
        hid_layout.addWidget(self.hid_entry)
        edit_layout.addLayout(hid_layout)
        
        # 字段选择
        col_layout = QHBoxLayout()
        col_layout.addWidget(QLabel("选择字段:"))
        self.col_combo = QComboBox()
        col_layout.addWidget(self.col_combo)
        edit_layout.addLayout(col_layout)
        
        # 新值
        val_layout = QHBoxLayout()
        val_layout.addWidget(QLabel("修正值:"))
        self.val_entry = QLineEdit()
        val_layout.addWidget(self.val_entry)
        edit_layout.addLayout(val_layout)
        
        self.update_btn = self.createStyledButton("💾 保存修改", "#e67e22", "#d35400")
        edit_layout.addWidget(self.update_btn)
        data_operations_layout.addWidget(edit_group)

    def createStatCard(self, parent_layout, title, value, color, row, col):
        card = QFrame()
        card.setStyleSheet(f"""
            QFrame {{
                background-color: {color};
                border-radius: 8px;
                padding: 10px;
                margin: 3px;
            }}
        """)
        card_layout = QVBoxLayout(card)
        card_layout.setContentsMargins(10, 10, 10, 10)
        
        title_label = QLabel(title)
        title_label.setFont(QFont("微软雅黑", 10, QFont.Bold))
        title_label.setStyleSheet("color: white;")
        title_label.setAlignment(Qt.AlignCenter)
        card_layout.addWidget(title_label)
        
        value_label = QLabel(value)
        value_label.setFont(QFont("Arial", 14, QFont.Bold))  # 改用 Arial 字体，更清晰
        value_label.setStyleSheet("color: white;")
        value_label.setAlignment(Qt.AlignCenter)
        card_layout.addWidget(value_label)
        
        parent_layout.addWidget(card, row, col)
        
        # 存储引用
        if not hasattr(self, 'stat_cards'):
            self.stat_cards = {}
        self.stat_cards[title] = value_label

    def createAnalysisTab(self):
        self.analysis_tab = QWidget()
        self.tab_widget.addTab(self.analysis_tab, "📊 深度分析报告")
        
        layout = QVBoxLayout(self.analysis_tab)
        
        self.analyze_btn = self.createStyledButton("📊 点击生成全维度统计报告", "#3498db", "#2980b9")
        layout.addWidget(self.analyze_btn)
        
        self.analysis_text = QTextEdit()
        self.analysis_text.setFont(QFont("Consolas", 11))
        layout.addWidget(self.analysis_text)

    def createVisualizationTab(self):
        self.viz_tab = QWidget()
        self.tab_widget.addTab(self.viz_tab, "📈 可视化与图表")
        
        layout = QVBoxLayout(self.viz_tab)
        
        # 图表选择区域
        chart_selector = QGroupBox("🎯 图表类型选择")
        chart_selector.setStyleSheet(self.getGroupBoxStyle())
        layout.addWidget(chart_selector)
        
        selector_layout = QHBoxLayout(chart_selector)
        
        self.income_btn = self.createStyledButton("💰 收入结构饼图", "#3498db", "#2980b9")
        self.land_btn = self.createStyledButton("🌾 土地分布直方图", "#2ecc71", "#27ae60")
        self.population_btn = self.createStyledButton("👥 人口结构分析", "#9b59b6", "#8e44ad")
        self.credit_btn = self.createStyledButton("💳 信贷获取情况", "#e74c3c", "#c0392b")
        self.pension_btn = self.createStyledButton("🏥 社保参保分析", "#f39c12", "#e67e22")
        self.export_chart_btn = self.createStyledButton("📈 导出图表", "#95a5a6", "#7f8c8d")
        
        selector_layout.addWidget(self.income_btn)
        selector_layout.addWidget(self.land_btn)
        selector_layout.addWidget(self.population_btn)
        selector_layout.addWidget(self.credit_btn)
        selector_layout.addWidget(self.pension_btn)
        selector_layout.addWidget(self.export_chart_btn)
        
        # 图表显示区域
        chart_group = QGroupBox("📈 图表展示区")
        chart_group.setStyleSheet(self.getGroupBoxStyle())
        layout.addWidget(chart_group)
        
        chart_layout = QVBoxLayout(chart_group)
        self.chart_frame = QWidget()
        chart_layout.addWidget(self.chart_frame)
        
        self.viz_status = QLabel("💡 请选择图表类型开始分析")
        self.viz_status.setStyleSheet("color: #7f8c8d; font-size: 12px;")
        chart_layout.addWidget(self.viz_status)

    def createQueryTab(self):
        self.query_tab = QWidget()
        self.tab_widget.addTab(self.query_tab, "🔍 自定义查询")
        
        layout = QVBoxLayout(self.query_tab)
        
        # 查询条件面板
        query_group = QGroupBox("📋 查询条件")
        query_group.setStyleSheet(self.getGroupBoxStyle())
        layout.addWidget(query_group)
        
        query_layout = QVBoxLayout(query_group)
        
        # 农户ID
        hid_layout = QHBoxLayout()
        hid_layout.addWidget(QLabel("🏠 农户ID (hid):"))
        self.query_hid_entry = QLineEdit()
        hid_layout.addWidget(self.query_hid_entry)
        query_layout.addLayout(hid_layout)
        
        # 按钮
        btn_layout = QHBoxLayout()
        self.filter_btn = self.createStyledButton("🔍 执行筛选", "#3498db", "#2980b9")
        self.export_report_btn = self.createStyledButton("📄 导出Word报表", "#2ecc71", "#27ae60")
        self.reset_btn = self.createStyledButton("🔄 重置查询", "#95a5a6", "#7f8c8d")
        
        btn_layout.addWidget(self.filter_btn)
        btn_layout.addWidget(self.export_report_btn)
        btn_layout.addWidget(self.reset_btn)
        query_layout.addLayout(btn_layout)
        
        # 结果显示
        result_group = QGroupBox("📊 查询结果")
        result_group.setStyleSheet(self.getGroupBoxStyle())
        layout.addWidget(result_group)
        
        result_layout = QVBoxLayout(result_group)
        self.query_result_text = QTextEdit()
        self.query_result_text.setReadOnly(True)
        result_layout.addWidget(self.query_result_text)

    def getGroupBoxStyle(self):
        return """
            QGroupBox {
                font-size: 14px;
                font-weight: bold;
                color: #2c3e50;
                border: 2px solid #e9ecef;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 10px 0 10px;
            }
        """

    def applyModernStyle(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f8f9fa;
            }
            QTabWidget::pane {
                border: 1px solid #e9ecef;
                background: white;
                border-radius: 5px;
            }
            QTabBar::tab {
                background: #ffffff;
                border: 1px solid #e9ecef;
                padding: 12px 20px;
                margin-right: 2px;
                border-radius: 5px 5px 0 0;
                color: #2c3e50;
                font-weight: bold;
            }
            QTabBar::tab:selected {
                background: #667eea;
                color: white;
            }
            QTabBar::tab:hover {
                background: #5a6fd8;
                color: white;
            }
            QLineEdit, QTextEdit, QComboBox {
                border: 1px solid #ced4da;
                border-radius: 4px;
                padding: 8px;
                background: white;
            }
            QLineEdit:focus, QTextEdit:focus {
                border-color: #667eea;
            }
        """)

    def connectSignals(self):
        self.browse_btn.clicked.connect(self.browse)
        self.load_btn.clicked.connect(self.load)
        self.add_btn.clicked.connect(self.add_row)
        self.update_btn.clicked.connect(self.update_cell)
        self.analyze_btn.clicked.connect(self.show_stats)
        self.income_btn.clicked.connect(lambda: self.plot_chart('income'))
        self.land_btn.clicked.connect(lambda: self.plot_chart('land'))
        self.population_btn.clicked.connect(lambda: self.plot_chart('population'))
        self.credit_btn.clicked.connect(lambda: self.plot_chart('credit'))
        self.pension_btn.clicked.connect(lambda: self.plot_chart('pension'))
        self.export_chart_btn.clicked.connect(self.export_chart)
        self.filter_btn.clicked.connect(self.apply_advanced_filter)
        self.export_report_btn.clicked.connect(self.export_report)
        self.reset_btn.clicked.connect(self.reset_query)

    # 保持原有的逻辑方法，但需要适配 PyQt5
    def browse(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择 Excel 文件", "", "Excel Files (*.xlsx)")
        if file_path:
            self.path_entry.setText(file_path)

    def load(self):
        p = self.path_entry.text()
        if not os.path.exists(p):
            QMessageBox.warning(self, "错误", "文件不存在")
            return
        try:
            self.manager = CLESDataManager(p)
            cols = list(self.manager.df.columns)
            self.col_combo.clear()
            self.col_combo.addItems(cols)
            
            # 更新字段显示
            fields_text = "\n".join([f"{i+1:03d} | {c}" for i, c in enumerate(cols)])
            self.fields_text.setPlainText(fields_text)
            
            # 更新统计卡片
            self.update_stats_cards()
            
            self.status_bar.showMessage(f"✅ 成功加载：{os.path.basename(p)} | 样本数：{len(self.manager.df)}")
            QMessageBox.information(self, "加载成功", f"数据已同步至管理系统\n样本数量：{len(self.manager.df)}")
        except Exception as e:
            QMessageBox.critical(self, "加载失败", f"文件格式错误或损坏：\n{str(e)}")

    def update_stats_cards(self):
        """更新统计卡片数据"""
        if not self.manager or not hasattr(self, 'stat_cards'):
            return
        
        df = self.manager.df
        
        # 样本总量
        sample_count = len(df)
        self.stat_cards["👥 样本总量"].setText(str(sample_count))
        
        # 村编码数
        village_count = df['a104'].nunique() if 'a104' in df.columns else 0
        self.stat_cards["🏘️ 村编码数"].setText(str(village_count))
        
        # 平均收入
        income_cols = ['b101', 'b102', 'b103', 'b104']
        available_income = [col for col in income_cols if col in df.columns]
        if available_income:
            total_income = df[available_income].sum(axis=1).mean()
            self.stat_cards["💰 平均收入"].setText(f"¥{total_income:,.0f}")
        else:
            self.stat_cards["💰 平均收入"].setText("¥0")
        
        # 平均耕地
        land_col = 'c107' if 'c107' in df.columns else ('c101' if 'c101' in df.columns else None)
        if land_col:
            avg_land = df[land_col].mean()
            self.stat_cards["🌾 平均耕地"].setText(f"{avg_land:.1f}亩")
        else:
            self.stat_cards["🌾 平均耕地"].setText("0亩")

    def add_row(self):
        if not self.manager:
            QMessageBox.warning(self, "提示", "请先加载数据")
            return
        raw = self.input_text.toPlainText().strip().split('\n')
        cols = self.manager.df.columns
        new_data = []
        for i in range(len(cols)):
            val = raw[i].strip() if i < len(raw) else ""
            if val == "":
                new_data.append(np.nan)
            else:
                try:
                    new_data.append(float(val) if '.' in val else int(val))
                except:
                    new_data.append(val)
        
        new_df = pd.DataFrame([new_data], columns=cols)
        self.manager.df = pd.concat([self.manager.df, new_df], ignore_index=True)
        if self.manager.save():
            QMessageBox.information(self, "成功", "新行数据已录入。")
            self.input_text.clear()
            # 更新统计卡片
            self.update_stats_cards()

    def update_cell(self):
        if not self.manager:
            QMessageBox.warning(self, "提示", "请先加载数据")
            return
        try:
            hid = int(self.hid_entry.text())
            col = self.col_combo.currentText()
            val = self.val_entry.text().strip()
            if val == "":
                val = np.nan
            else:
                try:
                    val = float(val) if '.' in val else int(val)
                except:
                    pass
            self.manager.df.loc[self.manager.df['hid'] == hid, col] = val
            if self.manager.save():
                QMessageBox.information(self, "成功", f"农户 {hid} 更新成功。")
                # 更新统计卡片
                self.update_stats_cards()
        except Exception as e:
            QMessageBox.critical(self, "修正失败", str(e))

    def show_stats(self):
        if not self.manager:
            QMessageBox.warning(self, "提示", "请先加载数据")
            return
        try:
            engine = CLESStatsEngine(self.manager.df)
            report = engine.run_full_analysis()
            self.analysis_text.setPlainText(report)
        except Exception as e:
            QMessageBox.critical(self, "分析失败", f"生成报告时出错：\n{str(e)}")
            print(f"Error in show_stats: {e}")  # 调试信息

    def plot_chart(self, chart_type):
        if not self.manager:
            QMessageBox.warning(self, "提示", "请先加载数据")
            return
        
        # 清空之前的图表和布局
        if self.chart_frame.layout():
            # 删除现有的布局
            old_layout = self.chart_frame.layout()
            while old_layout.count():
                child = old_layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()
            old_layout.deleteLater()
        
        # 关闭所有matplotlib figures
        plt.close('all')
        
        try:
            if chart_type == 'income':
                self.plot_income_structure()
                self.viz_status.setText("✅ 收入结构分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'land':
                self.plot_land_distribution()
                self.viz_status.setText("✅ 土地分布分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'population':
                self.plot_population_structure()
                self.viz_status.setText("✅ 人口结构分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'credit':
                self.plot_credit_analysis()
                self.viz_status.setText("✅ 信贷获取分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'pension':
                self.plot_pension_analysis()
                self.viz_status.setText("✅ 社保参保分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
        except Exception as e:
            self.viz_status.setText(f"❌ 图表生成失败: {str(e)}")
            self.viz_status.setStyleSheet("color: #e74c3c;")

    def plot_income_structure(self):
        inc_map = {'工资性': 'b101', '经营性': 'b102', '财产性': 'b103', '转移性': 'b104'}
        data = {k: self.manager.df[v].sum() for k, v in inc_map.items() if v in self.manager.df.columns}
        
        if not data:
            return
        
        plt.figure(figsize=(8, 6))
        plt.pie(data.values(), labels=data.keys(), autopct='%1.1f%%', startangle=140)
        plt.title("样本户均收入结构占比")
        
        canvas = FigureCanvas(plt.gcf())
        # 确保布局存在
        if not self.chart_frame.layout():
            layout = QVBoxLayout(self.chart_frame)
        else:
            layout = self.chart_frame.layout()
        layout.addWidget(canvas)

    def plot_land_distribution(self):
        if 'c107' not in self.manager.df.columns:
            QMessageBox.warning(self, "数据缺失", "土地面积数据列 'c107' 不存在，无法生成土地分布图表")
            return
        
        plt.figure(figsize=(8, 6))
        self.manager.df['c107'].hist(bins=20, color='#2ecc71', edgecolor='black', alpha=0.7)
        plt.title("农户经营土地面积分布 (亩)")
        plt.xlabel("面积")
        plt.ylabel("户数")
        plt.grid(True, alpha=0.3)
        
        canvas = FigureCanvas(plt.gcf())
        # 确保布局存在
        if not self.chart_frame.layout():
            layout = QVBoxLayout(self.chart_frame)
        else:
            layout = self.chart_frame.layout()
        layout.addWidget(canvas)

    def plot_population_structure(self):
        if 'a101' not in self.manager.df.columns:
            return
        
        plt.figure(figsize=(8, 6))
        plt.hist(self.manager.df['a101'], bins=15, color='#9b59b6', edgecolor='black', alpha=0.7)
        plt.title('农户家庭人口数分布')
        plt.xlabel('家庭人口数（人）')
        plt.ylabel('农户数量')
        plt.grid(True, alpha=0.3)
        
        canvas = FigureCanvas(plt.gcf())
        # 确保布局存在
        if not self.chart_frame.layout():
            layout = QVBoxLayout(self.chart_frame)
        else:
            layout = self.chart_frame.layout()
        layout.addWidget(canvas)

    def plot_credit_analysis(self):
        if 'f101' not in self.manager.df.columns:
            return
        
        credit_data = self.manager.df['f101'].value_counts()
        labels = ['有信贷' if k == 1 else '无信贷' for k in credit_data.index]
        
        plt.figure(figsize=(8, 6))
        plt.pie(credit_data.values, labels=labels, autopct='%1.1f%%', 
               colors=['#2ecc71', '#e74c3c'], startangle=90)
        plt.title('农户信贷获取情况分布')
        
        canvas = FigureCanvas(plt.gcf())
        # 确保布局存在
        if not self.chart_frame.layout():
            layout = QVBoxLayout(self.chart_frame)
        else:
            layout = self.chart_frame.layout()
        layout.addWidget(canvas)

    def plot_pension_analysis(self):
        if 'g101' not in self.manager.df.columns:
            QMessageBox.warning(self, "数据缺失", "社保参保数据列 'g101' 不存在，无法生成社保参保分析图表")
            return
        
        pension_data = self.manager.df['g101'].value_counts()
        labels = ['已参保' if k == 1 else '未参保' for k in pension_data.index]
        
        plt.figure(figsize=(8, 6))
        plt.bar(labels, pension_data.values, color=['#f39c12', '#95a5a6'])
        plt.title('城乡居民养老保险参保情况')
        plt.ylabel('农户数量')
        plt.grid(True, alpha=0.3)
        
        canvas = FigureCanvas(plt.gcf())
        # 确保布局存在
        if not self.chart_frame.layout():
            layout = QVBoxLayout(self.chart_frame)
        else:
            layout = self.chart_frame.layout()
        layout.addWidget(canvas)

    def export_chart(self):
        if not self.chart_frame.layout() or self.chart_frame.layout().count() == 0:
            QMessageBox.warning(self, "提示", "请先生成图表")
            return
        
        file_path, _ = QFileDialog.getSaveFileName(self, "保存图表", "", "PNG Files (*.png);;JPEG Files (*.jpg);;PDF Files (*.pdf)")
        if file_path:
            try:
                plt.savefig(file_path, dpi=300, bbox_inches='tight')
                QMessageBox.information(self, "成功", f"图表已保存至: {file_path}")
            except Exception as e:
                QMessageBox.critical(self, "保存失败", str(e))

    def apply_advanced_filter(self):
        if not self.manager:
            QMessageBox.warning(self, "提示", "请先加载数据")
            return
        
        hid_query = self.query_hid_entry.text().strip()
        
        if not hid_query:
            QMessageBox.warning(self, "提示", "请输入农户ID")
            return
        
        try:
            filtered_df = self.manager.df.copy()
            
            if hid_query:
                hid_val = int(hid_query)
                filtered_df = filtered_df[filtered_df['hid'] == hid_val]
            
            if len(filtered_df) == 0:
                QMessageBox.information(self, "查询结果", "未找到匹配的记录")
                self.query_result_text.setPlainText("❌ 未找到匹配的记录")
                return
            
            # 更新统计信息
            # 获取各项数据
            total_income = filtered_df[['b101', 'b102', 'b103', 'b104']].sum(axis=1).sum()
            wage_income = filtered_df['b101'].sum() if 'b101' in filtered_df.columns else 0
            business_income = filtered_df['b102'].sum() if 'b102' in filtered_df.columns else 0
            property_income = filtered_df['b103'].sum() if 'b103' in filtered_df.columns else 0
            transfer_income = filtered_df['b104'].sum() if 'b104' in filtered_df.columns else 0
            
            land_area = filtered_df.get('c107', pd.Series([0]*len(filtered_df))).sum()
            family_size = filtered_df.get('a101', pd.Series([0]*len(filtered_df))).sum()
            
            credit_status = "有信贷" if filtered_df.get('f101', pd.Series([0]*len(filtered_df))).sum() > 0 else "无信贷"
            pension_status = "已参保" if filtered_df.get('g101', pd.Series([0]*len(filtered_df))).sum() > 0 else "未参保"
            
            stats_text = f"""✅ 查询成功！
📊 匹配记录数: {len(filtered_df)} 条

🏠 农户ID: {hid_val}

👥 家庭人口数: {family_size} 人

💰 收入情况:
   • 总收入: ¥{total_income:,.0f}
   • 工资性收入: ¥{wage_income:,.0f}
   • 经营性收入: ¥{business_income:,.0f}
   • 财产性收入: ¥{property_income:,.0f}
   • 转移性收入: ¥{transfer_income:,.0f}

🌾 农业情况:
   • 经营土地面积: {land_area:.2f} 亩

🏦 金融社保:
   • 信贷获取: {credit_status}
   • 养老保险: {pension_status}"""
            
            self.query_result_text.setPlainText(stats_text)
            self.status_bar.showMessage(f"查询完成 - 找到 {len(filtered_df)} 条记录")
            
            # 临时显示筛选结果
            self.manager.df = filtered_df
            
        except ValueError:
            QMessageBox.critical(self, "错误", "请输入有效的数字ID")
        except Exception as e:
            QMessageBox.critical(self, "查询失败", str(e))

    def reset_query(self):
        self.query_hid_entry.clear()
        self.query_result_text.setPlainText("💡 请先加载数据并设置查询条件")
        self.status_bar.showMessage("就绪 - 等待查询操作")
        
        if hasattr(self, 'original_df'):
            self.manager.df = self.original_df.copy()
        QMessageBox.information(self, "重置完成", "查询条件已清空")

    def export_report(self):
        if not self.manager:
            QMessageBox.warning(self, "提示", "请先加载数据")
            return
        rep = CLESReportGenerator(self.manager)
        rep.export_word_report()
        main_paned = tk.PanedWindow(self.tab_a, orient="horizontal", bg="#f5f6fa", sashwidth=4)
        main_paned.pack(fill="both", expand=True)

        # 数据概览面板
        overview_frame = tk.Frame(main_paned, bg="white", relief="solid", bd=1)
        main_paned.add(overview_frame, width=400)
        
        overview_title = tk.Label(overview_frame, text="📈 数据概览", 
                                 font=("微软雅黑", 14, "bold"), bg="white", fg="#2c3e50")
        overview_title.pack(pady=(20, 15))
        
        # 统计卡片网格
        stats_grid = tk.Frame(overview_frame, bg="white")
        stats_grid.pack(padx=20, pady=(0, 20))
        
        # 创建统计卡片
        self.create_stats_card(stats_grid, "👥 样本总量", "0", "#3498db", 0, 0)
        self.create_stats_card(stats_grid, "🏘️ 村编码数", "0", "#2ecc71", 0, 1)
        self.create_stats_card(stats_grid, "💰 平均收入", "¥0", "#9b59b6", 1, 0)
        self.create_stats_card(stats_grid, "🌾 平均耕地", "0亩", "#e74c3c", 1, 1)
        
        # 字段参照面板
        fields_panel = tk.Frame(main_paned, bg="white", relief="solid", bd=1)
        main_paned.add(fields_panel, width=300)
        
        fields_title = tk.Label(fields_panel, text="📋 字段对照表", 
                               font=("微软雅黑", 14, "bold"), bg="white", fg="#2c3e50")
        fields_title.pack(pady=(20, 15))
        
        self.fields_display = scrolledtext.ScrolledText(fields_panel, 
                                                       font=("Consolas", 10), 
                                                       bg="#f8f9fa", fg="#2f3640", 
                                                       borderwidth=1, relief="solid", 
                                                       padx=15, pady=15)
        self.fields_display.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        self.fields_display.config(state="disabled")

        # 2. 数据输入 (中)
        f_input = tk.LabelFrame(main_paned, text=" ➕ 手动录入新行 ", font=("微软雅黑", 10, "bold"), 
                               bg="white", padx=10, pady=10, relief="flat")
        main_paned.add(f_input, width=600)
        self.input_box = scrolledtext.ScrolledText(f_input, font=("Consolas", 11), bg="white", 
                                                  fg="#2f3640", padx=10, pady=10)
        self.input_box.pack(fill="both", expand=True)
        
        btn_add = tk.Button(f_input, text=" 📥 提交并同步文件 ", command=self.add_row, 
                           font=("微软雅黑", 10, "bold"), bg="#2ecc71", fg="white", borderwidth=0, pady=8, cursor="hand2")
        btn_add.pack(fill="x", pady=10)

        # 3. 精确修改 (右)
        f_edit = tk.LabelFrame(main_paned, text=" 🔧 单元格修正 ", font=("微软雅黑", 10, "bold"), 
                              bg="white", padx=15, pady=10, relief="flat")
        main_paned.add(f_edit, width=300)
        
        def create_label_entry(parent, text):
            tk.Label(parent, text=text, font=("微软雅黑", 9), bg="white", fg="#7f8c8d").pack(anchor="w", pady=(10,0))
            ent = tk.Entry(parent, font=("微软雅黑", 10), bg="#f5f6fa", borderwidth=0)
            ent.pack(fill="x", pady=5, ipady=3)
            return ent

        self.ent_hid = create_label_entry(f_edit, "农户 HID")
        tk.Label(f_edit, text="选择字段", font=("微软雅黑", 9), bg="white", fg="#7f8c8d").pack(anchor="w", pady=(10,0))
        self.cb_col = ttk.Combobox(f_edit, state="readonly", font=("微软雅黑", 10))
        self.cb_col.pack(fill="x", pady=5)
        self.ent_val = create_label_entry(f_edit, "修正值")
        
        tk.Button(f_edit, text=" 💾 保存修改 ", command=self.update_cell, 
                  font=("微软雅黑", 10, "bold"), bg="#e67e22", fg="white", borderwidth=0, pady=8, cursor="hand2").pack(fill="x", pady=30)

    def init_stats_ui(self):
        f = tk.Frame(self.tab_b, bg="#f5f6fa", padx=20, pady=20)
        f.pack(fill="both", expand=True)
        
        btn_run = tk.Button(f, text=" 📊 点击生成全维度统计报告 ", command=self.show_stats, 
                           font=("微软雅黑", 11, "bold"), bg="#3498db", fg="white", borderwidth=0, padx=20, pady=10, cursor="hand2")
        btn_run.pack(anchor="w", pady=(0, 15))
        
        self.stats_box = scrolledtext.ScrolledText(f, font=("Consolas", 11), bg="#ffffff", 
                                                  fg="#2c3e50", borderwidth=1, relief="solid", padx=20, pady=20)
        self.stats_box.pack(fill="both", expand=True)

    # --- 逻辑控制 ---
    def browse(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择 Excel 文件", "", "Excel Files (*.xlsx)")
        if file_path:
            self.path_entry.setText(file_path)

    def load(self):
        p = self.path_entry.text()
        if not os.path.exists(p):
            QMessageBox.warning(self, "错误", "文件不存在")
            return
        try:
            self.manager = CLESDataManager(p)
            cols = list(self.manager.df.columns)
            self.col_combo.clear()
            self.col_combo.addItems(cols)
            
            # 更新字段显示
            fields_text = "\n".join([f"{i+1:03d} | {c}" for i, c in enumerate(cols)])
            self.fields_text.setPlainText(fields_text)
            
            # 更新统计卡片
            self.update_stats_cards()
            
            self.status_bar.showMessage(f"✅ 成功加载：{os.path.basename(p)} | 样本数：{len(self.manager.df)}")
            QMessageBox.information(self, "加载成功", f"数据已同步至管理系统\n样本数量：{len(self.manager.df)}")
        except Exception as e:
            QMessageBox.critical(self, "加载失败", f"文件格式错误或损坏：\n{str(e)}")

    def add_row(self):
        if not self.manager: return
        raw = self.input_box.get("1.0", tk.END).strip().split('\n')
        cols = self.manager.df.columns
        new_data = []
        for i in range(len(cols)):
            val = raw[i].strip() if i < len(raw) else ""
            if val == "": new_data.append(np.nan)
            else:
                try: new_data.append(float(val) if '.' in val else int(val))
                except: new_data.append(val)
        
        new_df = pd.DataFrame([new_data], columns=cols)
        self.manager.df = pd.concat([self.manager.df, new_df], ignore_index=True)
        if self.manager.save():
            messagebox.showinfo("成功", "新行数据已录入。")
            self.input_box.delete("1.0", tk.END)
            # 更新统计卡片
            self.update_stats_cards()

    def update_cell(self):
        if not self.manager: return
        try:
            hid = int(self.hid_entry.text())
            col = self.col_combo.currentText()
            val = self.val_entry.text().strip()
            if val == "": val = np.nan
            else:
                try: val = float(val) if '.' in val else int(val)
                except: pass
            self.manager.df.loc[self.manager.df['hid'] == hid, col] = val
            if self.manager.save():
                QMessageBox.information(self, "成功", f"农户 {hid} 更新成功。")
                # 更新统计卡片
                self.update_stats_cards()
        except Exception as e:
            QMessageBox.critical(self, "修正失败", str(e))

    def plot_chart(self, chart_type):
        if not self.manager:
            QMessageBox.warning(self, "提示", "请先加载数据")
            return
        
        # 清空之前的图表和布局
        if self.chart_frame.layout():
            # 删除现有的布局
            old_layout = self.chart_frame.layout()
            while old_layout.count():
                child = old_layout.takeAt(0)
                if child.widget():
                    child.widget().deleteLater()
            old_layout.deleteLater()
        
        # 关闭所有matplotlib figures
        plt.close('all')
        
        try:
            if chart_type == 'income':
                self.plot_income_structure()
                self.viz_status.setText("✅ 收入结构分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'land':
                self.plot_land_distribution()
                self.viz_status.setText("✅ 土地分布分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'population':
                self.plot_population_structure()
                self.viz_status.setText("✅ 人口结构分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'credit':
                self.plot_credit_analysis()
                self.viz_status.setText("✅ 信贷获取分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
            elif chart_type == 'pension':
                self.plot_pension_analysis()
                self.viz_status.setText("✅ 社保参保分析图表已生成")
                self.viz_status.setStyleSheet("color: #27ae60;")
        except Exception as e:
            self.viz_status.setText(f"❌ 图表生成失败: {str(e)}")
            self.viz_status.setStyleSheet("color: #e74c3c;")

    def plot_population_structure(self):
        if 'a101' not in self.manager.df.columns:
            return
        
        plt.figure(figsize=(8, 6))
        plt.hist(self.manager.df['a101'], bins=15, color='#9b59b6', edgecolor='black', alpha=0.7)
        plt.title('农户家庭人口数分布')
        plt.xlabel('家庭人口数（人）')
        plt.ylabel('农户数量')
        plt.grid(True, alpha=0.3)
        
        canvas = FigureCanvas(plt.gcf())
        # 确保布局存在
        if not self.chart_frame.layout():
            layout = QVBoxLayout(self.chart_frame)
        else:
            layout = self.chart_frame.layout()
        layout.addWidget(canvas)

    def plot_credit_analysis(self):
        if 'f101' not in self.manager.df.columns:
            return
        
        credit_data = self.manager.df['f101'].value_counts()
        labels = ['有信贷' if k == 1 else '无信贷' for k in credit_data.index]
        
        plt.figure(figsize=(8, 6))
        plt.pie(credit_data.values, labels=labels, autopct='%1.1f%%', 
               colors=['#2ecc71', '#e74c3c'], startangle=90)
        plt.title('农户信贷获取情况分布')
        
        canvas = FigureCanvas(plt.gcf())
        # 确保布局存在
        if not self.chart_frame.layout():
            layout = QVBoxLayout(self.chart_frame)
        else:
            layout = self.chart_frame.layout()
        layout.addWidget(canvas)

    def export_chart(self):
        """导出当前图表"""
        if not hasattr(self, 'chart_frame_tab_c') or not self.chart_frame_tab_c.winfo_children():
            messagebox.showwarning("提示", "请先生成图表")
            return
        
        file_path = filedialog.asksaveasfilename(defaultextension=".png",
                                                filetypes=[("PNG files", "*.png"),
                                                          ("JPEG files", "*.jpg"),
                                                          ("PDF files", "*.pdf")])
        if file_path:
            try:
                plt.savefig(file_path, dpi=300, bbox_inches='tight')
                messagebox.showinfo("成功", f"图表已保存至: {file_path}")
            except Exception as e:
                messagebox.showerror("保存失败", str(e))

    def init_custom_query_ui(self):
        # 主容器 - 使用渐变背景
        self.tab_d_frame = tk.Frame(self.tab_d, bg="#667eea")
        self.tab_d_frame.pack(fill="both", expand=True, padx=0, pady=0)
        
        # 创建渐变背景画布
        self.create_gradient_background(self.tab_d_frame)
        
        # 标题区域
        title_frame = tk.Frame(self.tab_d_frame, bg="#667eea")
        title_frame.pack(fill="x", pady=(20, 10))
        
        title_label = tk.Label(title_frame, text="🔍 高级查询与分析中心", 
                              font=("微软雅黑", 18, "bold"), bg="#667eea", fg="white")
        title_label.pack()
        
        subtitle_label = tk.Label(title_frame, text="智能筛选 · 数据洞察 · 精准导出", 
                                 font=("微软雅黑", 10), bg="#667eea", fg="#e8f4f8")
        subtitle_label.pack(pady=(5, 0))
        
        # 主内容区域
        main_content = tk.Frame(self.tab_d_frame, bg="#f8f9fa")
        main_content.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # 左侧查询面板
        query_panel = tk.Frame(main_content, bg="white", relief="solid", bd=1)
        query_panel.pack(side="left", fill="y", padx=(0, 15), pady=15)
        
        # 查询面板标题
        panel_title = tk.Label(query_panel, text="📋 查询条件", 
                              font=("微软雅黑", 14, "bold"), bg="white", fg="#2c3e50")
        panel_title.pack(pady=(20, 15))
        
        # 农户ID查询
        hid_frame = tk.Frame(query_panel, bg="white")
        hid_frame.pack(fill="x", padx=20, pady=(0, 15))
        
        hid_label = tk.Label(hid_frame, text="🏠 农户ID (hid):", 
                            font=("微软雅黑", 11, "bold"), bg="white", fg="#34495e")
        hid_label.pack(anchor="w", pady=(0, 5))
        
        self.ent_hid_query = tk.Entry(hid_frame, font=("微软雅黑", 11), 
                                     bg="#f8f9fa", relief="solid", bd=1, 
                                     insertbackground="#3498db", width=25)
        self.ent_hid_query.pack(fill="x", ipady=8, padx=(0, 10))
        
        # 村编码查询
        village_frame = tk.Frame(query_panel, bg="white")
        village_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        village_label = tk.Label(village_frame, text="🏘️ 村编码 (a104):", 
                                font=("微软雅黑", 11, "bold"), bg="white", fg="#34495e")
        village_label.pack(anchor="w", pady=(0, 5))
        
        self.ent_village_query = tk.Entry(village_frame, font=("微软雅黑", 11), 
                                         bg="#f8f9fa", relief="solid", bd=1, 
                                         insertbackground="#3498db", width=25)
        self.ent_village_query.pack(fill="x", ipady=8, padx=(0, 10))
        
        # 操作按钮区域
        button_frame = tk.Frame(query_panel, bg="white")
        button_frame.pack(fill="x", padx=20, pady=(0, 20))
        
        # 筛选按钮
        filter_btn = tk.Button(button_frame, text="🔍 执行筛选", 
                              command=self.apply_advanced_filter,
                              font=("微软雅黑", 11, "bold"), bg="#3498db", fg="white",
                              relief="flat", bd=0, padx=25, pady=12, cursor="hand2",
                              activebackground="#2980b9", activeforeground="white")
        filter_btn.pack(fill="x", pady=(0, 10))
        
        # 导出按钮
        export_btn = tk.Button(button_frame, text="📄 导出Word报表", 
                              command=self.export_report,
                              font=("微软雅黑", 11, "bold"), bg="#2ecc71", fg="white",
                              relief="flat", bd=0, padx=25, pady=12, cursor="hand2",
                              activebackground="#27ae60", activeforeground="white")
        export_btn.pack(fill="x", pady=(0, 10))
        
        # 重置按钮
        reset_btn = tk.Button(button_frame, text="🔄 重置查询", 
                             command=self.reset_query,
                             font=("微软雅黑", 11, "bold"), bg="#95a5a6", fg="white",
                             relief="flat", bd=0, padx=25, pady=12, cursor="hand2",
                             activebackground="#7f8c8d", activeforeground="white")
        reset_btn.pack(fill="x")
        
        # 右侧结果显示区域
        result_panel = tk.Frame(main_content, bg="white", relief="solid", bd=1)
        result_panel.pack(side="right", fill="both", expand=True, pady=15)
        
        # 结果面板标题
        result_title = tk.Label(result_panel, text="📊 查询结果与可视化", 
                               font=("微软雅黑", 14, "bold"), bg="white", fg="#2c3e50")
        result_title.pack(pady=(20, 15))
        
        # 统计信息显示
        stats_frame = tk.Frame(result_panel, bg="#f8f9fa", relief="solid", bd=1)
        stats_frame.pack(fill="x", padx=20, pady=(0, 15))
        
        self.stats_label = tk.Label(stats_frame, 
                                   text="💡 请先加载数据并设置查询条件", 
                                   font=("微软雅黑", 10), bg="#f8f9fa", fg="#7f8c8d",
                                   justify="left", anchor="w")
        self.stats_label.pack(fill="x", padx=15, pady=15)
        
        # 图表显示区域
        self.chart_frame_tab_d = tk.Frame(result_panel, bg="white")
        self.chart_frame_tab_d.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        
        # 底部状态栏
        status_frame = tk.Frame(self.tab_d_frame, bg="#34495e", height=30)
        status_frame.pack(fill="x", side="bottom")
        status_frame.pack_propagate(False)
        
        self.query_status = tk.Label(status_frame, text="就绪 - 等待查询操作", 
                                    font=("微软雅黑", 9), bg="#34495e", fg="#ecf0f1")
        self.query_status.pack(side="left", padx=20)
        
        # 添加悬停效果
        self.add_hover_effects(filter_btn, "#2980b9")
        self.add_hover_effects(export_btn, "#27ae60")
        self.add_hover_effects(reset_btn, "#7f8c8d")

    def create_gradient_background(self, parent):
        """创建渐变背景"""
        canvas = tk.Canvas(parent, highlightthickness=0)
        canvas.pack(fill="both", expand=True)
        
        # 创建渐变效果
        width = 800  # 默认宽度
        height = 600  # 默认高度
        
        # 绘制渐变背景
        for i in range(height):
            # 从深蓝到浅蓝的渐变
            r = int(102 + (255 - 102) * (i / height) * 0.3)
            g = int(126 + (255 - 126) * (i / height) * 0.3)
            b = int(234 + (255 - 234) * (i / height) * 0.3)
            color = f'#{r:02x}{g:02x}{b:02x}'
            canvas.create_line(0, i, width, i, fill=color)
        
        # 将其他组件放在画布上
        canvas.create_text(width//2, height//2 - 100, text="🌟", 
                          font=("Arial", 60), fill="#ffffff", anchor="center")
        canvas.create_text(width//2, height//2 + 50, text="高级查询系统", 
                          font=("微软雅黑", 24, "bold"), fill="#ffffff", anchor="center")

    def add_hover_effects(self, button, hover_color):
        """添加按钮悬停效果"""
        def on_enter(e):
            button.config(bg=hover_color)
        def on_leave(e):
            button.config(bg=button.cget("bg"))
        button.bind("<Enter>", on_enter)
        button.bind("<Leave>", on_leave)

    def export_report(self):
        """导出Word报表"""
        if not self.manager:
            messagebox.showwarning("提示", "请先加载数据")
            return
        rep = CLESReportGenerator(self.manager)
        rep.export_word_report()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = CLESApp()
    window.show()
    sys.exit(app.exec_())